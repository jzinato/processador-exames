import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import matplotlib.pyplot as plt
import datetime
import base64
from io import StringIO
import plotly.express as px
import plotly.graph_objects as go
from PIL import Image
# Adicione estas linhas ao início do arquivo, junto com os outros imports
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuração da página
st.set_page_config(
    page_title="Processador de Exames Médicos",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Título e descrição
st.title("Processador de Exames Médicos")
st.markdown("Ferramenta para processamento de PDFs de exames médicos")

# Inicializar variáveis de estado da sessão
if 'exam_history' not in st.session_state:
    st.session_state.exam_history = []
    
if 'current_exam' not in st.session_state:
    st.session_state.current_exam = None
    
if 'patient_info' not in st.session_state:
    st.session_state.patient_info = {"name": "", "collectionDate": ""}

if 'selected_metrics' not in st.session_state:
    st.session_state.selected_metrics = []

# Função para calcular eGFR usando CKD-EPI
def calculate_ckd_epi(creatinine, age, is_female=False, is_black=False):
    cr = float(creatinine)
    
    if is_female:
        if cr <= 0.7:
            egfr = 144 * (cr / 0.7) ** -0.329 * 0.993 ** age
        else:
            egfr = 144 * (cr / 0.7) ** -1.209 * 0.993 ** age
    else:
        if cr <= 0.9:
            egfr = 141 * (cr / 0.9) ** -0.411 * 0.993 ** age
        else:
            egfr = 141 * (cr / 0.9) ** -1.209 * 0.993 ** age
    
    # Ajuste para raça negra
    if is_black:
        egfr *= 1.159
    
    return round(egfr)

# Verificar se um resultado está fora do intervalo de referência
def is_abnormal(value, reference):
    if not reference:
        return False
    
    # Extrair números do valor de referência
    ref_match = re.findall(r'[\d,.]+', reference)
    if not ref_match or len(ref_match) < 2:
        return False
    
    min_ref = float(ref_match[0].replace(',', '.'))
    max_ref = float(ref_match[1].replace(',', '.'))
    
    # Extrair o valor numérico do resultado
    val_match = re.search(r'[\d,.]+', str(value))
    if not val_match:
        return False
    
    val = float(val_match.group().replace(',', '.'))
    
    return val < min_ref or val > max_ref

# Processar texto do PDF
def process_pdf_text(text):
    try:
        # Estrutura para armazenar os resultados
        categories = {
            'Hemograma': [],
            'Bioquímica': [],
            'Hormonais': [],
            'Outros': [],
            'Imagem': []
        }
        
        # Extrair informações do paciente
        name_match = re.search(r'Nome:\s*(.*)', text)
        date_match = re.search(r'Data da Coleta:\s*(.*)', text)
        
        patient_info = {
            "name": name_match.group(1).strip() if name_match else "",
            "collectionDate": date_match.group(1).strip() if date_match else ""
        }
        
        # Processar cada linha do texto
        lines = text.split('\n')
        current_category = 'Outros'
        in_image_section = False
        
        for line in lines:
            line = line.strip()
            
            # Verificar seções principais
            if 'EXAMES DE IMAGEM:' in line:
                in_image_section = True
                continue
            
            # Identificar categorias com base em palavras-chave
            if 'Hemograma:' in line:
                current_category = 'Hemograma'
                continue
            elif any(keyword in line for keyword in [
                'Proteínas Totais', 'Bilirrubinas:', 'Ureia:', 'Creatinina:',
                'Cálcio:', 'Potássio:', 'Fósforo:', 'Bicarbonato:',
                'Ferro Sérico:', 'Fosfatase Alcalina:'
            ]):
                current_category = 'Bioquímica'
            elif any(keyword in line for keyword in [
                'Testosterona', 'PSA', 'Paratormônio'
            ]):
                current_category = 'Hormonais'
            
            # Processar resultados
            result_match = re.search(r'○\s*(.*?):\s*(.*?)\s*\(Referência:\s*(.*?)\)', line)
            simple_result_match = re.search(r'●\s*(.*?):\s*(.*?)\s*\(Referência:\s*(.*?)\)', line)
            
            if result_match or simple_result_match:
                match = result_match or simple_result_match
                name = match.group(1).strip()
                value = match.group(2).strip()
                reference = match.group(3).strip()
                
                # Extrair valor numérico para gráficos
                numeric_value = None
                if value:
                    num_match = re.search(r'[\d,.]+', value)
                    if num_match:
                        numeric_value = float(num_match.group().replace(',', '.'))
                
                unit = re.sub(r'[\d,.\s]+', '', value).strip() if value else ""
                
                categories[current_category].append({
                    "name": name,
                    "value": value,
                    "reference": reference,
                    "numericValue": numeric_value,
                    "unit": unit,
                    "isAbnormal": is_abnormal(value, reference)
                })
            elif in_image_section:
                # Processar resultados de imagem
                image_finding_match = re.search(r'○\s*(.*?):\s*(.*)', line)
                if image_finding_match:
                    categories['Imagem'].append({
                        "name": image_finding_match.group(1).strip(),
                        "value": image_finding_match.group(2).strip(),
                        "isAbnormal": True  # Consideramos todos os achados de imagem como relevantes
                    })
        
        # Verificar se precisamos calcular o clearance de creatinina
        has_creatinine = any(item["name"] == "Creatinina" for item in categories['Bioquímica'])
        has_egfr = any(item["name"] == "Estimativa do Ritmo de Filtração Glomerular" for item in categories['Bioquímica'])
        
        if has_creatinine and not has_egfr:
            creatinine_item = next((item for item in categories['Bioquímica'] if item["name"] == "Creatinina"), None)
            if creatinine_item:
                cr_value = float(re.search(r'[\d,.]+', creatinine_item["value"]).group().replace(',', '.'))
                # Estimando idade como 65 para demonstração
                estimated_age = 65
                calculated_egfr = calculate_ckd_epi(cr_value, estimated_age, False, False)
                
                categories['Bioquímica'].append({
                    "name": "Estimativa do Ritmo de Filtração Glomerular (CKD-EPI)",
                    "value": f"{calculated_egfr} mL/min/1,73m²",
                    "numericValue": calculated_egfr,
                    "unit": "mL/min/1,73m²",
                    "reference": "> 90 mL/min/1,73m²",
                    "isAbnormal": calculated_egfr < 90,
                    "isCalculated": True
                })
        
        return patient_info, categories
    
    except Exception as e:
        st.error(f"Erro ao processar o arquivo: {str(e)}")
        return None, None

# Função para exibir resultados de exames com estilo
def display_exam_results(exam_data, show_title=True):
    if show_title:
        st.subheader(f"Relatório de Exames")
    
    # Hemograma
    if exam_data['Hemograma']:
        st.markdown("### Hemograma")
        cols = st.columns(2)
        for i, item in enumerate(exam_data['Hemograma']):
            col = cols[i % 2]
            with col:
                if item['isAbnormal']:
                    st.markdown(f"""
                    <div style="background-color: rgba(255, 0, 0, 0.1); border-left: 4px solid #EF5350; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']}</span>
                            <span style="font-weight: 700; color: #D32F2F;">{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div style="background-color: #F5F5F5; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']}</span>
                            <span>{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
    
    # Bioquímica
    if exam_data['Bioquímica']:
        st.markdown("### Bioquímica")
        cols = st.columns(2)
        for i, item in enumerate(exam_data['Bioquímica']):
            col = cols[i % 2]
            with col:
                if item.get('isCalculated', False):
                    st.markdown(f"""
                    <div style="background-color: rgba(33, 150, 243, 0.1); border-left: 4px solid #2196F3; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']} <span style="font-size: 0.8em; color: #2196F3;">(calculado)</span></span>
                            <span style="font-weight: 700; color: {'#D32F2F' if item['isAbnormal'] else '#000'};">{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                elif item['isAbnormal']:
                    st.markdown(f"""
                    <div style="background-color: rgba(255, 0, 0, 0.1); border-left: 4px solid #EF5350; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']}</span>
                            <span style="font-weight: 700; color: #D32F2F;">{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div style="background-color: #F5F5F5; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']}</span>
                            <span>{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
    
    # Hormonais
    if exam_data['Hormonais']:
        st.markdown("### Hormonais")
        cols = st.columns(2)
        for i, item in enumerate(exam_data['Hormonais']):
            col = cols[i % 2]
            with col:
                if item['isAbnormal']:
                    st.markdown(f"""
                    <div style="background-color: rgba(255, 0, 0, 0.1); border-left: 4px solid #EF5350; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']}</span>
                            <span style="font-weight: 700; color: #D32F2F;">{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div style="background-color: #F5F5F5; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']}</span>
                            <span>{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
    
    # Outros
    if exam_data['Outros']:
        st.markdown("### Outros Exames")
        cols = st.columns(2)
        for i, item in enumerate(exam_data['Outros']):
            col = cols[i % 2]
            with col:
                if item['isAbnormal']:
                    st.markdown(f"""
                    <div style="background-color: rgba(255, 0, 0, 0.1); border-left: 4px solid #EF5350; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']}</span>
                            <span style="font-weight: 700; color: #D32F2F;">{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div style="background-color: #F5F5F5; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <div style="display: flex; justify-content: space-between;">
                            <span style="font-weight: 500;">{item['name']}</span>
                            <span>{item['value']}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
    
    # Imagem
    if exam_data['Imagem']:
        st.markdown("### Exames de Imagem")
        with st.container():
            st.markdown("""
            <div style="background-color: #F5F5F5; padding: 15px; border-radius: 5px;">
            """, unsafe_allow_html=True)
            
            for item in exam_data['Imagem']:
                st.markdown(f"""
                <div style="margin-bottom: 8px;">
                    <span style="font-weight: 500;">{item['name']}:</span>
                    <span style="margin-left: 8px;">{item['value']}</span>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)

# Função para adicionar exames históricos de exemplo
def load_sample_data():
    # Exame mais recente
    sample_report_text = """Resultados de Exames para PEP
Dados do Paciente:
● Nome: Nicomedes Ferreira Filho
● Data da Coleta: 17/02/2025
Resultados:
● Hemograma:
○ Hemoglobina: 11,4 g/dL (Referência: 13,0 a 17,0 g/dL)
○ VCM: 91,0 fL (Referência: 83,0 a 101,0 fL)
○ HCM: 30,3 pg (Referência: 27,0 a 32,0 pg)
○ Leucócitos: 5340/µL (Referência: 4000 a 10000/µL)
○ Segmentados: 56,9% (Referência: 2000,0 a 7000,0)
○ Eosinófilos: 13,5% (Referência: 20,0 a 500,0)
○ Linfócitos: 19,3% (Referência: 1000,0 a 3000,0)
○ Monócitos: 8,4% (Referência: 200,0 a 1000,0)
○ Basófilos: 1,9% (Referência: 20,0 a 100,0)
○ Plaquetas: 192.000/µL (Referência: 150 a 400 mil/µL)
● Ferro Sérico: 98 µg/dL (Referência: Mulheres: 50-170 µg/dL, Homens: 65-175 µg/dL)
● Capacidade Total de Combinação do Ferro: 259 µg/dL (Referência: 250 a 450 µg/dL)
● Ferritina: 106 ng/mL (Referência: Homens: 21,81-274,66 ng/mL, Mulheres: 4,63-204,00 ng/mL)
● Proteínas Totais e Fracionadas:
○ Proteínas Totais: 6,6 g/dL (Referência: 6,4-8,3 g/dL)
○ Albumina: 4,2 g/dL (Referência: 3,5-5,0 g/dL)
○ Globulina: 2,4 g/dL
○ Relação A/G: 1,8
● Bilirrubinas:
○ Total: 0,55 mg/dL (Referência: Adulto: 0,2 a 1,2 mg/dL)
○ Direta: 0,26 mg/dL (Referência: Adulto: 0,0 a 0,5 mg/dL)
○ Indireta: 0,29 mg/dL (Referência: Adulto: 0,2 a 0,7 mg/dL)
● Fosfatase Alcalina: 140 U/L (Referência: 22 a 79 anos, Homens: 50-116 U/L)
● Ureia: 145,00 mg/dL (Referência: Adultos: 12,8-42,8 mg/dL, Adultos > 60 anos: 17,1-49,2 mg/dL)
● Creatinina: 3,18 mg/dL (Referência: Adultos: 0,5-1,00 mg/dL, Homem > 60 anos: 0,6-1,20 mg/dL)
● Cálcio: 8,8 mg/dL (Referência: Adulto: 8,4 a 10,2 mg/dL, Homem > 60 anos: 8,8 a 10,0 mg/dL)
● Estimativa do Ritmo de Filtração Glomerular: 18 mL/min/1,73m^2 (Referência: Adultos > 18 anos: > 90 mL/min/1,73 m^2)
● Potássio: 4,50 mmol/L (Referência: 3,5 a 5,1 mmol/L)
● Cálcio Iônico: 1,11 mEq/L (Referência: 1,16 a 1,32 mEq/L)
● Fósforo: 4,3 mg/dL (Referência: Adultos: 2,5-4,5 mg/dL)
● Paratormônio PTH Intacto (Molécula Inteira): 265,1 pg/mL (Referência: 15 a 68,3 pg/mL)
● Bicarbonato: 16 mEq/L (Referência: 20 a 32 mEq/L)
● Testosterona Total: 12,49 ng/dL (Referência: Homens > 50 anos: 220,91 a 715,81 ng/dL)
● PSA Total e Livre:
○ PSA Livre: 0,14 ng/mL (Referência: 0,0 a 0,5 ng/mL)
○ PSA Total: 0,54 ng/mL (Referência: 0,0 a 4,0 ng/mL)
○ Porcentagem de PSA Livre/PSA Total: 26%
● Testosterona Livre Calculada: 0,15 ng/dL (Referência: Homem de 50 a 89 anos: 1,81 a 10,20 ng/dL)
2. EXAMES DE IMAGEM:
Ultrassonografia dos Rins e Vias Urinárias:
○ Rins: Sinais de nefropatia parenquimatosa crônica bilateral. Pequenos cistos renais simples corticais bilaterais.
○ Bexiga: Pós-miccional de 66,7 mL. Sinais de bexiga."""

    patient_info, exam_data = process_pdf_text(sample_report_text)
    
    st.session_state.patient_info = patient_info
    st.session_state.current_exam = exam_data
    
    # Adicionar exames históricos de demonstração
    historical_exams = [
        {
            "date": "17/01/2025",
            "patient_info": {"name": "Nicomedes Ferreira Filho", "collectionDate": "17/01/2025"},
            "data": {
                'Hemograma': [
                    {"name": "Hemoglobina", "value": "11,2 g/dL", "reference": "13,0 a 17,0 g/dL", "isAbnormal": True, "numericValue": 11.2, "unit": "g/dL"},
                    {"name": "Leucócitos", "value": "5420/µL", "reference": "4000 a 10000/µL", "isAbnormal": False, "numericValue": 5420, "unit": "/µL"},
                    {"name": "Plaquetas", "value": "188.000/µL", "reference": "150 a 400 mil/µL", "isAbnormal": False, "numericValue": 188000, "unit": "/µL"}
                ],
                'Bioquímica': [
                    {"name": "Ureia", "value": "142,00 mg/dL", "reference": "Adultos: 12,8-42,8 mg/dL", "isAbnormal": True, "numericValue": 142.0, "unit": "mg/dL"},
                    {"name": "Creatinina", "value": "3,22 mg/dL", "reference": "Adultos: 0,5-1,00 mg/dL", "isAbnormal": True, "numericValue": 3.22, "unit": "mg/dL"},
                    {"name": "Estimativa do Ritmo de Filtração Glomerular", "value": "17 mL/min/1,73m²", "reference": "Adultos > 18 anos: > 90 mL/min/1,73 m²", "isAbnormal": True, "numericValue": 17, "unit": "mL/min/1,73m²"}
                ],
                'Hormonais': [],
                'Outros': [],
                'Imagem': []
            }
        },
        {
            "date": "17/12/2024",
            "patient_info": {"name": "Nicomedes Ferreira Filho", "collectionDate": "17/12/2024"},
            "data": {
                'Hemograma': [
                    {"name": "Hemoglobina", "value": "10,8 g/dL", "reference": "13,0 a 17,0 g/dL", "isAbnormal": True, "numericValue": 10.8, "unit": "g/dL"},
                    {"name": "Leucócitos", "value": "5500/µL", "reference": "4000 a 10000/µL", "isAbnormal": False, "numericValue": 5500, "unit": "/µL"},
                    {"name": "Plaquetas", "value": "180.000/µL", "reference": "150 a 400 mil/µL", "isAbnormal": False, "numericValue": 180000, "unit": "/µL"}
                ],
                'Bioquímica': [
                    {"name": "Ureia", "value": "138,00 mg/dL", "reference": "Adultos: 12,8-42,8 mg/dL", "isAbnormal": True, "numericValue": 138.0, "unit": "mg/dL"},
                    {"name": "Creatinina", "value": "3,34 mg/dL", "reference": "Adultos: 0,5-1,00 mg/dL", "isAbnormal": True, "numericValue": 3.34, "unit": "mg/dL"},
                    {"name": "Estimativa do Ritmo de Filtração Glomerular", "value": "16 mL/min/1,73m²", "reference": "Adultos > 18 anos: > 90 mL/min/1,73 m²", "isAbnormal": True, "numericValue": 16, "unit": "mL/min/1,73m²"}
                ],
                'Hormonais': [],
                'Outros': [],
                'Imagem': []
            }
        },
        {
            "date": "17/10/2024",
            "patient_info": {"name": "Nicomedes Ferreira Filho", "collectionDate": "17/10/2024"},
            "data": {
                'Hemograma': [
                    {"name": "Hemoglobina", "value": "10,5 g/dL", "reference": "13,0 a 17,0 g/dL", "isAbnormal": True, "numericValue": 10.5, "unit": "g/dL"},
                    {"name": "Leucócitos", "value": "5600/µL", "reference": "4000 a 10000/µL", "isAbnormal": False, "numericValue": 5600, "unit": "/µL"},
                    {"name": "Plaquetas", "value": "175.000/µL", "reference": "150 a 400 mil/µL", "isAbnormal": False, "numericValue": 175000, "unit": "/µL"}
                ],
                'Bioquímica': [
                    {"name": "Ureia", "value": "132,00 mg/dL", "reference": "Adultos: 12,8-42,8 mg/dL", "isAbnormal": True, "numericValue": 132.0, "unit": "mg/dL"},
                    {"name": "Creatinina", "value": "3,40 mg/dL", "reference": "Adultos: 0,5-1,00 mg/dL", "isAbnormal": True, "numericValue": 3.40, "unit": "mg/dL"},
                    {"name": "Estimativa do Ritmo de Filtração Glomerular", "value": "15 mL/min/1,73m²", "reference": "Adultos > 18 anos: > 90 mL/min/1,73 m²", "isAbnormal": True, "numericValue": 15, "unit": "mL/min/1,73m²"}
                ],
                'Hormonais': [],
                'Outros': [],
                'Imagem': []
            }
        }
    ]
    
    st.session_state.exam_history = historical_exams
    
    # Selecionar algumas métricas importantes por padrão
    default_metrics = [
        {"category": "Hemograma", "name": "Hemoglobina", "unit": "g/dL"},
        {"category": "Bioquímica", "name": "Creatinina", "unit": "mg/dL"},
        {"category": "Bioquímica", "name": "Ureia", "unit": "mg/dL"},
        {"category": "Bioquímica", "name": "Estimativa do Ritmo de Filtração Glomerular", "unit": "mL/min/1,73m²"}
    ]
    
    st.session_state.selected_metrics = default_metrics

# Função para exibir o histórico de exames
def show_exam_history():
    if not st.session_state.exam_history:
        st.info("Nenhum histórico de exames disponível.")
        return
    
    st.subheader("Histórico de Exames")
    
    for i, exam in enumerate(st.session_state.exam_history):
        with st.expander(f"Exame de {exam['date']}", expanded=(i == 0)):
            display_exam_results(exam['data'], show_title=False)

# Função para exibir gráficos de tendência
def show_graphs():
    if not st.session_state.exam_history and not st.session_state.current_exam:
        st.info("Nenhum dado disponível para gráficos.")
        return
    
    st.subheader("Gráficos de Evolução")
    
    # Preparar lista de métricas disponíveis
    available_metrics = []
    
    if st.session_state.current_exam:
        for category in st.session_state.current_exam:
            for item in st.session_state.current_exam[category]:
                if item.get('numericValue') is not None:
                    # Verificar se tem dados históricos
                    has_history = False
                    for exam in st.session_state.exam_history:
                        if category in exam['data']:
                            for hist_item in exam['data'][category]:
                                if hist_item['name'] == item['name'] and hist_item.get('numericValue') is not None:
                                    has_history = True
                                    break
                    
                    if has_history or category == 'Bioquímica':  # Sempre mostrar itens de bioquímica
                        metric = {
                            "category": category,
                            "name": item['name'],
                            "unit": item.get('unit', '')
                        }
                        if metric not in available_metrics:
                            available_metrics.append(metric)
    
    # Interface para selecionar métricas
    st.markdown("### Selecione Parâmetros")
    
    metric_buttons = []
    for metric in available_metrics:
        is_selected = any(m['name'] == metric['name'] and m['category'] == metric['category'] 
                       for m in st.session_state.selected_metrics)
        metric_buttons.append({"metric": metric, "selected": is_selected})
    
    # Organizar botões em linhas
    cols = st.columns(3)
    for i, button_data in enumerate(metric_buttons):
        metric = button_data["metric"]
        is_selected = button_data["selected"]
        
        col = cols[i % 3]
        with col:
            if st.button(
                f"{metric['name']}",
                key=f"metric_{metric['category']}_{metric['name']}",
                type="primary" if is_selected else "secondary",
                use_container_width=True
            ):
                # Alternar seleção
                if is_selected:
                    st.session_state.selected_metrics = [
                        m for m in st.session_state.selected_metrics 
                        if not (m['name'] == metric['name'] and m['category'] == metric['category'])
                    ]
                else:
                    st.session_state.selected_metrics.append(metric)
                
                # Atualização UI
                st.rerun()
    
    # Função para preparar dados para os gráficos
    def prepare_graph_data(metric_name, metric_category):
        data = []
        
        # Dados históricos
        for exam in st.session_state.exam_history:
            if metric_category in exam['data']:
                item = next((item for item in exam['data'][metric_category] if item['name'] == metric_name), None)
                if item and item.get('numericValue') is not None:
                    data.append({
                        "date": exam['date'],
                        "value": item['numericValue']
                    })
        
        # Dados atuais
        if st.session_state.current_exam and metric_category in st.session_state.current_exam:
            current_item = next((item for item in st.session_state.current_exam[metric_category] 
                             if item['name'] == metric_name), None)
            if current_item and current_item.get('numericValue') is not None:
                # Verificar se já existe um ponto com esta data
                if not any(point['date'] == st.session_state.patient_info['collectionDate'] for point in data):
                    data.append({
                        "date": st.session_state.patient_info['collectionDate'],
                        "value": current_item['numericValue']
                    })
        
        # Ordenar por data
        data.sort(key=lambda x: datetime.datetime.strptime(x['date'], '%d/%m/%Y'))
        
        return data
    
    # Exibir gráficos para métricas selecionadas
    if st.session_state.selected_metrics:
        for metric in st.session_state.selected_metrics:
            data = prepare_graph_data(metric['name'], metric['category'])
            
            if len(data) < 2:
                st.warning(f"Dados insuficientes para gerar gráfico de {metric['name']} (mínimo 2 pontos).")
                continue
            
            # Criar DataFrame para plotly
            df = pd.DataFrame(data)
            
            # Criar gráfico com plotly
            fig = px.line(
                df, 
                x='date', 
                y='value', 
                markers=True,
                title=f"{metric['name']} ({metric['unit']})",
                template="simple_white"
            )
            
            fig.update_layout(
                xaxis_title="Data",
                yaxis_title=f"Valor ({metric['unit']})",
                height=400,
                margin=dict(l=20, r=20, t=40, b=20),
            )
            
            st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Selecione parâmetros acima para visualizar gráficos.")
# Função para gerar documento Word
def generate_word_report(patient_info, exam_data):
    # Criar um novo documento
    doc = Document()
    
    # Configurar estilo do documento
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Adicionar cabeçalho
    header = doc.add_heading('RELATÓRIO DE EXAMES MÉDICOS', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações do paciente
    doc.add_paragraph()
    doc.add_paragraph(f"Paciente: {patient_info['name']}")
    doc.add_paragraph(f"Data da coleta: {patient_info['collectionDate']}")
    doc.add_paragraph()
    
    # Adicionar linha horizontal
    doc.add_paragraph().add_run('_' * 80).bold = True
    
    # Função para adicionar uma seção de exames
    def add_exam_section(title, items):
        if not items:
            return
            
        doc.add_heading(title, level=2)
        
        # Adicionar tabela
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Cabeçalhos da tabela
        header_cells = table.rows[0].cells
        header_cells[0].text = "Exame"
        header_cells[1].text = "Resultado"
        
        # Formatar cabeçalhos
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            
        # Adicionar resultados
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item['name']
            
            # Adicionar valor com formatação para anormais
            result_paragraph = row_cells[1].paragraphs[0]
            result_run = result_paragraph.add_run(item['value'])
            
            if item.get('isAbnormal', False):
                result_run.bold = True
                result_run.font.color.rgb = RGBColor(192, 0, 0)  # Vermelho
                
            # Adicionar nota se for calculado
            if item.get('isCalculated', False):
                calc_run = result_paragraph.add_run(" (calculado)")
                calc_run.italic = True
                calc_run.font.size = Pt(9)
    
        doc.add_paragraph()
    
    # Adicionar cada seção de exames
    add_exam_section("Hemograma", exam_data['Hemograma'])
    add_exam_section("Bioquímica", exam_data['Bioquímica'])
    add_exam_section("Hormonais", exam_data['Hormonais'])
    add_exam_section("Outros Exames", exam_data['Outros'])
    
    # Adicionar exames de imagem
    if exam_data['Imagem']:
        doc.add_heading("Exames de Imagem", level=2)
        for item in exam_data['Imagem']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{item['name']}: ").bold = True
            p.add_run(item['value'])
    
    # Adicionar rodapé
    doc.add_paragraph()
    footer = doc.add_paragraph("Relatório gerado automaticamente pelo Processador de Exames Médicos")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento em memória
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    return docx_stream
# Função para gerar PDF para download
def create_download_link(content, filename):
    b64 = base64.b64encode(content.encode()).decode()
    href = f'<a href="data:file/txt;base64,{b64}" download="{filename}">Baixar Relatório</a>'
    return href

# Interface principal
with st.sidebar:
    st.header("Opções")
    
    # Upload de arquivo
    uploaded_file = st.file_uploader("Faça upload do PDF do exame", type=["pdf", "txt"])
    
    # Botão para carregar dados de exemplo
    if st.button("Carregar Dados de Exemplo"):
        load_sample_data()
    
    # Informações do paciente (se disponíveis)
    if st.session_state.patient_info and st.session_state.patient_info["name"]:
        st.markdown("---")
        st.subheader("Dados do Paciente")
        st.markdown(f"**Nome:** {st.session_state.patient_info['name']}")
        st.markdown(f"**Data da Coleta:** {st.session_state.patient_info['collectionDate']}")

# Processar arquivo enviado
# Processar arquivo enviado
if uploaded_file is not None:
    try:
        # Verificar tipo de arquivo
        if uploaded_file.name.endswith('.pdf'):
            # Processar arquivo PDF
            import PyPDF2
            from io import BytesIO
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.getvalue()))
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text()
        else:
            # Processar arquivo de texto
            text_content = StringIO(uploaded_file.getvalue().decode("utf-8")).read()
        
        patient_info, exam_data = process_pdf_text(text_content)
        
        # Atualizar o estado da sessão
        st.session_state.patient_info = patient_info
        st.session_state.current_exam = exam_data
        
        # Adicionar ao histórico se for um novo exame
        if patient_info["collectionDate"] and not any(exam["date"] == patient_info["collectionDate"] for exam in st.session_state.exam_history):
            new_exam = {
                "date": patient_info["collectionDate"],
                "patient_info": patient_info,
                "data": exam_data
            }
            
            # Inserir mantendo ordem cronológica (mais recente primeiro)
            st.session_state.exam_history = sorted(
                st.session_state.exam_history + [new_exam],
                key=lambda x: datetime.datetime.strptime(x["date"], "%d/%m/%Y"),
                reverse=True
            )
        
        st.success("Arquivo processado com sucesso!")
        # Remover esta linha: st.rerun()
        
    except Exception as e:
        st.error(f"Erro ao processar o arquivo: {str(e)}")

# Criar abas
tabs = st.tabs(["Exame Atual", "Histórico", "Gráficos"])

# Aba de Exame Atual
with tabs[0]:
    if st.session_state.current_exam:
        # Botões para exportação
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Exportar como Texto", key="export_text"):
                report_text = f"RELATÓRIO DE EXAMES\n\nPaciente: {st.session_state.patient_info['name']}\nData: {st.session_state.patient_info['collectionDate']}\n\n"
                # Código para gerar conteúdo texto do relatório
                st.markdown(create_download_link(report_text, "relatorio_exame.txt"), unsafe_allow_html=True)
                
        with col2:
            if st.button("Exportar como Word", key="export_word"):
                # Gerar documento Word
                docx_file = generate_word_report(st.session_state.patient_info, st.session_state.current_exam)
                
                # Converter para base64 para download
                b64 = base64.b64encode(docx_file.getvalue()).decode()
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="relatorio_exame_{st.session_state.patient_info["collectionDate"].replace("/", "-")}.docx">Baixar Relatório Word</a>'
                st.markdown(href, unsafe_allow_html=True)
        
        display_exam_results(st.session_state.current_exam)
    else:
        st.info("Carregue um arquivo de exame ou use os dados de exemplo para visualizar os resultados.")

# Aba de Histórico
with tabs[1]:
    show_exam_history()

# Aba de Gráficos
with tabs[2]:
    show_graphs()

# Adicionar CSS customizado
st.markdown("""
<style>
    .st-emotion-cache-1y4p8pa {padding-top: 2rem !important;}
    .block-container {padding-top: 2rem !important;}
    
    /* Customização dos botões de métrica */
    .stButton button {
        text-overflow: ellipsis;
        overflow: hidden;
        white-space: nowrap;
    }
</style>
""", unsafe_allow_html=True)